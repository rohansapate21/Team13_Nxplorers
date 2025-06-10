import os
import docx
import spacy
import nltk
import pandas as pd
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Add all possible NLTK data paths
nltk.data.path.append(os.path.expanduser('~/AppData/Roaming/nltk_data'))
nltk.data.path.append(os.path.join(os.path.dirname(os.path.abspath(__file__)), '../../nltk_data'))
nltk.data.path.append('/usr/share/nltk_data')
nltk.data.path.append('/usr/local/share/nltk_data')
nltk.data.path.append('/usr/lib/nltk_data')
nltk.data.path.append('/usr/local/lib/nltk_data')

# Ensure NLTK and spaCy resources are available
try:
    stop_words = set(stopwords.words('english'))
except LookupError:
    nltk.download('stopwords')
    stop_words = set(stopwords.words('english'))
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')
try:
    nlp = spacy.load('en_core_web_sm')
except OSError:
    from spacy.cli import download
    download('en_core_web_sm')
    nlp = spacy.load('en_core_web_sm')

print("NLTK data paths:", nltk.data.path)

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return '\n'.join([para.text for para in doc.paragraphs])

def extract_text_from_file(file):
    # file: Django InMemoryUploadedFile or path
    if hasattr(file, 'temporary_file_path'):
        return extract_text_from_docx(file.temporary_file_path())
    elif hasattr(file, 'name') and file.name.endswith('.docx'):
        file.seek(0)
        doc = docx.Document(file)
        return '\n'.join([para.text for para in doc.paragraphs])
    else:
        return ''

def preprocess(text):
    text = text.lower()
    tokens = word_tokenize(text)
    tokens = [word for word in tokens if word.isalpha() and word not in stop_words]
    return ' '.join(tokens)

def extract_skills(text):
    doc = nlp(text)
    skills = []
    for chunk in doc.noun_chunks:
        if chunk.root.pos_ in ['NOUN', 'PROPN']:
            skills.append(chunk.text.lower())
    return list(set(skills))

def get_cosine_similarity(text1, text2):
    vectorizer = TfidfVectorizer()
    tfidf = vectorizer.fit_transform([text1, text2])
    return cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0] * 100

def analyze_resume_vs_jd(resume_file, jd_text):
    # resume_file: Django InMemoryUploadedFile
    resume_text = extract_text_from_file(resume_file)
    resume_clean = preprocess(resume_text)
    jd_clean = preprocess(jd_text)
    resume_skills = extract_skills(resume_text)
    jd_skills = extract_skills(jd_text)
    match_score = get_cosine_similarity(resume_clean, jd_clean)
    matched_skills = list(set(resume_skills) & set(jd_skills))
    missing_skills = list(set(jd_skills) - set(resume_skills))
    report = {
        "match_score": round(match_score, 2),
        "resume_skills": resume_skills,
        "jd_skills": jd_skills,
        "matched_skills": matched_skills,
        "missing_skills": missing_skills,
        "resume_text": resume_text,
        "jd_text": jd_text,
    }
    return report

class ResumeParser:
    def __init__(self):
        self.nlp = spacy.load('en_core_web_sm')
        self.stop_words = set(stopwords.words('english'))

    def parse_resume(self, resume_file):
        """Parse a resume file and extract relevant information."""
        resume_text = extract_text_from_file(resume_file)
        if not resume_text:
            raise ValueError("Could not extract text from resume file")

        # Extract basic information
        doc = self.nlp(resume_text)
        
        # Extract skills
        skills = extract_skills(resume_text)
        
        # Extract education (simple pattern matching)
        education = []
        education_keywords = ['bachelor', 'master', 'phd', 'degree', 'university', 'college']
        for sent in doc.sents:
            if any(keyword in sent.text.lower() for keyword in education_keywords):
                education.append(sent.text.strip())
        
        # Extract experience (simple pattern matching)
        experience = []
        experience_keywords = ['experience', 'worked', 'job', 'position', 'role']
        for sent in doc.sents:
            if any(keyword in sent.text.lower() for keyword in experience_keywords):
                experience.append(sent.text.strip())
        
        return {
            'text': resume_text,
            'skills': skills,
            'education': education,
            'experience': experience
        } 